import streamlit as st
from docx import Document
from docx.shared import Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image as PILImage
import io
from datetime import datetime
import re

st.set_page_config(page_title="照片報告生成器", page_icon="📸", layout="wide")

st.title("📸 照片報告生成器 (文字輸入版)")
st.markdown("**Report Date & Delivery Date：直接輸入 DD/MM/YYYY 格式**")

# 初始化
if 'photos' not in st.session_state:
    st.session_state.photos = []
if 'descriptions' not in st.session_state:
    st.session_state.descriptions = []

# 文字輸入日期（DD/MM/YYYY）
col1, col2 = st.columns(2)
with col1:
    report_date = st.text_input("📅 Report Date", value=datetime.now().strftime("%d/%m/%Y"), 
                               placeholder="02/03/2026", help="格式：日/月/年 如 02/03/2026")
with col2:
    delivery_date = st.text_input("📦 Delivery Date", value=datetime.now().strftime("%d/%m/%Y"), 
                                 placeholder="02/03/2026", help="格式：日/月/年 如 02/03/2026")

# 驗證日期格式
def is_valid_date(date_str):
    pattern = r'^\d{1,2}/\d{1,2}/\d{4}$'
    return re.match(pattern, date_str) is not None

if report_date and not is_valid_date(report_date):
    st.error("❌ Report Date 格式錯誤！請用 DD/MM/YYYY (如 02/03/2026)")
if delivery_date and not is_valid_date(delivery_date):
    st.error("❌ Delivery Date 格式錯誤！請用 DD/MM/YYYY (如 02/03/2026)")

# 🔧 修正：移除 key="photos"
uploaded_files = st.file_uploader("選擇照片", type=['jpg','jpeg','png','webp'], 
                                 accept_multiple_files=True)

# 處理照片上傳
if uploaded_files:
    remaining = 8 - len(st.session_state.photos)
    for file in uploaded_files[:remaining]:
        if len(st.session_state.photos) < 8:
            img = PILImage.open(file)
            st.session_state.photos.append(img)
            st.session_state.descriptions.append("")
    st.success(f"✅ 已新增 {len(uploaded_files)} 張照片")

# 照片編輯區
if st.session_state.photos:
    st.subheader(f"📷 已選擇 {len(st.session_state.photos)}/8 張")
    
    cols = st.columns(2)
    for i, (photo, desc) in enumerate(zip(st.session_state.photos, st.session_state.descriptions)):
        with cols[i % 2]:
            st.image(photo, caption=f"照片 {i+1}", use_column_width=True)
            new_desc = st.text_input(f"照片 {i+1} 描述", value=desc, key=f"desc_{i}")
            st.session_state.descriptions[i] = new_desc
    
    col1, col2 = st.columns([1, 2])
    with col1:
        if st.button("🗑️ 清除全部", type="secondary"):
            st.session_state.photos = []
            st.session_state.descriptions = []
            st.rerun()
    with col2:
        if (st.button("📄 生成報告", type="primary") and 
            is_valid_date(report_date) and 
            is_valid_date(delivery_date)):
            generate_report(report_date, delivery_date)
        elif st.button("📄 生成報告", type="primary"):
            st.warning("⚠️ 請修正日期格式後再生成！")

else:
    st.info("👆 上傳照片開始使用")

# 生成報告函數
def generate_report(report_date, delivery_date):
    """生成包含文字日期的報告"""
    doc = Document()
    
    # 標題
    title = doc.add_heading('照片報告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 日期資訊（直接使用文字輸入）
    date_section = doc.add_paragraph()
    date_section.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    date_run1 = date_section.add_run(f"Report Date: {report_date}\n")
    date_run1.bold = True
    date_run1.font.size = Cm(0.6)
    
    date_run2 = date_section.add_run(f"Delivery Date: {delivery_date}")
    date_run2.bold = True
    date_run2.font.size = Cm(0.6)
    
    # 照片表格
    doc.add_paragraph("")  # 空行
    p = doc.add_paragraph("照片列表：")
    p.bold = True
    
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    
    for i in range(min(8, len(st.session_state.photos))):
        cell = table.rows[i//2].cells[i%2]
        
        # 圖片
        img_copy = st.session_state.photos[i].copy()
        img_copy.thumbnail((2*Inches, 2*Inches), PILImage.Resampling.LANCZOS)
        img_buffer = io.BytesIO()
        img_copy.save(img_buffer, 'JPEG', quality=90)
        
        run = cell.paragraphs[0].add_run()
        run.add_picture(img_buffer, width=Cm(4))
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 描述
        desc_para = cell.add_paragraph(st.session_state.descriptions[i] or '無描述')
        desc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 頁尾
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_para.add_run(f"生成時間: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    
    # 儲存檔案
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    st.success(f"✅ 報告生成完成！\n📅 Report Date: {report_date}\n📦 Delivery Date: {delivery_date}")
    
    st.download_button(
        label="📥 下載 Word 報告",
        data=buffer.getvalue(),
        file_name=f"照片報告_{report_date}_{delivery_date}_{timestamp}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True
    )

st.markdown("---")
st.markdown("*文字輸入日期 • DD/MM/YYYY格式 • Powered by Streamlit*")
